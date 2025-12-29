#!/usr/bin/env python3
# -*- coding: utf-8 -*-

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Read Russian text
    with open('ОТВЕТ_APPLE_24_12_2025_ОЧИЩЕННЫЙ_RU.txt', 'r', encoding='utf-8') as f:
        ru_text = f.read()
    
    # Read English text
    with open('ОТВЕТ_APPLE_24_12_2025_ОЧИЩЕННЫЙ_EN.txt', 'r', encoding='utf-8') as f:
        en_text = f.read()
    
    # Create document
    doc = Document()
    
    # Add Russian section
    doc.add_heading('ОТВЕТ APPLE - РУССКИЙ', 0)
    for line in ru_text.split('\n'):
        if line.strip():
            if line.startswith('Framework:') or line.startswith('Файл:') or line.startswith('File:') or line.startswith('Extension:') or line.startswith('Использование:') or line.startswith('Usage:') or line.startswith('Реализация:') or line.startswith('Implementation:') or line.startswith('Разрешение:') or line.startswith('Permission:') or line.startswith('Endpoints:'):
                p = doc.add_paragraph(line.strip())
                p.style = 'List Bullet'
            elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')) and not line.strip()[2:3].isalpha():
                p = doc.add_paragraph(line.strip())
                p.style = 'List Number'
            elif line.strip().startswith('-'):
                p = doc.add_paragraph(line.strip())
                p.style = 'List Bullet'
            elif len(line.strip()) > 0 and line.strip()[0].isupper() and not line.strip().startswith('Framework') and not line.strip().startswith('Файл') and not line.strip().startswith('File'):
                doc.add_heading(line.strip(), level=1)
            else:
                doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph()
    
    # Add page break
    doc.add_page_break()
    
    # Add English section
    doc.add_heading('APPLE RESPONSE - ENGLISH', 0)
    for line in en_text.split('\n'):
        if line.strip():
            if line.startswith('Framework:') or line.startswith('File:') or line.startswith('Extension:') or line.startswith('Usage:') or line.startswith('Implementation:') or line.startswith('Permission:') or line.startswith('Endpoints:'):
                p = doc.add_paragraph(line.strip())
                p.style = 'List Bullet'
            elif line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')) and not line.strip()[2:3].isalpha():
                p = doc.add_paragraph(line.strip())
                p.style = 'List Number'
            elif line.strip().startswith('-'):
                p = doc.add_paragraph(line.strip())
                p.style = 'List Bullet'
            elif len(line.strip()) > 0 and line.strip()[0].isupper() and not line.strip().startswith('Framework') and not line.strip().startswith('File'):
                doc.add_heading(line.strip(), level=1)
            else:
                doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph()
    
    # Save document
    doc.save('ОТВЕТ_APPLE_24_12_2025_RU_EN.docx')
    print("DOCX file created successfully!")
    
except ImportError:
    print("python-docx library not installed. Installing...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    print("Please run the script again.")
except Exception as e:
    print(f"Error: {e}")

