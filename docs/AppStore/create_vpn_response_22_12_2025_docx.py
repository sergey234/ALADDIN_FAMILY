#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Создает DOCX файл с ответом Apple (22 декабря 2025)
Сначала английский текст, потом русский
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Создаем новый документ
doc = Document()

# АНГЛИЙСКИЙ ТЕКСТ
doc.add_heading('Response to Apple Review Team - Guideline 5.4', 0)

p = doc.add_paragraph()
p.add_run('Dear App Review Team,').bold = True

doc.add_paragraph()
doc.add_paragraph('Thank you for your feedback. We have conducted a complete review of our code and confirm:')

p = doc.add_paragraph()
p.add_run('ALADDIN IS NOT A VPN APPLICATION').bold = True
p.add_run('\n\nALADDIN is a family security and parental control application. It does not provide VPN functionality.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('IMPORTANT:').bold = True

doc.add_paragraph('• We DO NOT have the ability to switch to a corporate developer account.')
doc.add_paragraph('• We DO NOT use VPN technologies in the application.')
doc.add_paragraph('• We are ready to fix whatever you tell us! Please tell us what and where to fix? We are ready to provide any evidence of the absence of VPN in the code. Please tell us what we need to do so that the ALADDIN application does not resemble a VPN? In the very first build on 03.12.25, VPN occupied only 2% of the functionality of the entire ALADDIN application, then we removed everything!')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('TECHNICAL EVIDENCE OF NO VPN:').bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('1. NetworkExtension framework IS NOT USED:').bold = True
doc.add_paragraph('   • Framework is not added to project.pbxproj')
doc.add_paragraph('   • All NetworkExtension imports are commented out and not executed')
doc.add_paragraph('   • 0 active NetworkExtension imports in the build code')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('2. NO VPN EXTENSION TARGET:').bold = True
doc.add_paragraph('   • ALADDINPacketTunnel target has been completely removed from project.pbxproj')
doc.add_paragraph('   • PacketTunnelProvider.swift has been removed')
doc.add_paragraph('   • No VPN extension in the project')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('3. NO VPN ENTITLEMENTS:').bold = True
doc.add_paragraph('   • All ALADDINPacketTunnel.entitlements files have been removed from the active project')
doc.add_paragraph('   • No com.apple.developer.networking.vpn.api in the project')
doc.add_paragraph('   • Check of project.pbxproj: 0 mentions of VPN entitlements')
doc.add_paragraph('   • No VPN servers in settings')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('4. VPN API CALLS ARE ABSENT:').bold = True
doc.add_paragraph('   • startVPNTunnel() - not executed (commented out and not compiled)')
doc.add_paragraph('   • stopVPNTunnel() - not executed (commented out and not compiled)')
doc.add_paragraph('   • All VPN API methods are commented out and not compiled')
doc.add_paragraph('   • 0 active VPN API calls in the compiled code')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('5. THE APPLICATION USES ONLY STANDARD iOS APIs:').bold = True
doc.add_paragraph('   • URLSession for network requests')
doc.add_paragraph('   • Standard security protocols')
doc.add_paragraph('   • No system extensions (network extensions)')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('6. CHECK OF ALL IPA FILE COMPONENTS:').bold = True
doc.add_paragraph('   • Binary code — VPN absent')
doc.add_paragraph('   • Resources — VPN absent')
doc.add_paragraph('   • Info.plist — VPN absent')
doc.add_paragraph('   • Entitlements — VPN absent (only in BACKUPS, inactive files)')
doc.add_paragraph('   • Frameworks — NetworkExtension absent')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('COMPLETE REMOVAL OF VPN').bold = True
doc.add_paragraph('All VPN-related components have been removed:')
doc.add_paragraph('• All VPN files and targets removed')
doc.add_paragraph('• All VPN classes and methods renamed')
doc.add_paragraph('• All VPN strings in the user interface updated')
doc.add_paragraph('• VPN entitlements and configurations removed')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('REQUEST TO APPLE TEAM:').bold = True
doc.add_paragraph('To quickly resolve the issue, please indicate to us where exactly you see VPN functionality:')
doc.add_paragraph('1. Screenshot of the application screen where VPN functionality is visible')
doc.add_paragraph('2. Or exact indication of the location in the code/metadata where you see VPN')
doc.add_paragraph('3. Or the name of a specific screen/function that looks like VPN')

doc.add_paragraph()
doc.add_paragraph('We have conducted a complete code audit and do not find any VPN functionality. All checks show that the code is clean of VPN components. We are ready to immediately fix any location you specify.')

doc.add_paragraph()
doc.add_paragraph('If the application for some reason looks like a VPN but is not one, we are ready to rename any interface elements or descriptions that may cause such an association.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Build 14 contains all fixes.').bold = True

doc.add_paragraph()
doc.add_paragraph('Best regards,')
doc.add_paragraph('Sergey Khlystov')
doc.add_paragraph('ALADDIN')
doc.add_paragraph('December 22, 2025')

# Разрыв страницы
doc.add_page_break()

# РУССКИЙ ТЕКСТ
doc.add_heading('Ответ команде App Review - Guideline 5.4', 0)

p = doc.add_paragraph()
p.add_run('Уважаемая команда App Review,').bold = True

doc.add_paragraph()
doc.add_paragraph('Благодарим вас за обратную связь. Мы провели полную проверку нашего кода и подтверждаем:')

p = doc.add_paragraph()
p.add_run('ALADDIN НЕ ЯВЛЯЕТСЯ VPN-ПРИЛОЖЕНИЕМ').bold = True
p.add_run('\n\nПриложение ALADDIN - это приложение для семейной безопасности и родительского контроля. Оно не предоставляет VPN-функциональность.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ВАЖНО:').bold = True

doc.add_paragraph('• У нас НЕТ возможности перехода на корпоративный аккаунт разработчика.')
doc.add_paragraph('• Мы НЕ используем VPN-технологии в приложении.')
doc.add_paragraph('• Мы готовы исправить, что вы нам скажите! Скажите пожалуйста что и где исправить? Готовы предоставить любые доказательства отсутствия VPN в коде. Скажите пожалуйста, что нам сделать что бы приложение ALADDIN не было похожим на VPN? В самой первой сборке 03.12.25 VPN занимало всего 2% функционала всего приложения ALADDIN, потом мы все удалили!')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ТЕХНИЧЕСКИЕ ДОКАЗАТЕЛЬСТВА ОТСУТСТВИЯ VPN:').bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('1. NetworkExtension framework НЕ ИСПОЛЬЗУЕТСЯ:').bold = True
doc.add_paragraph('   • Framework не добавлен в project.pbxproj')
doc.add_paragraph('   • Все импорты NetworkExtension закомментированы и не выполняются')
doc.add_paragraph('   • 0 активных импортов NetworkExtension в коде сборки')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('2. НЕТ VPN EXTENSION TARGET:').bold = True
doc.add_paragraph('   • ALADDINPacketTunnel target полностью удален из project.pbxproj')
doc.add_paragraph('   • PacketTunnelProvider.swift удален')
doc.add_paragraph('   • Нет VPN extension в проекте')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('3. НЕТ VPN ENTITLEMENTS:').bold = True
doc.add_paragraph('   • Все файлы ALADDINPacketTunnel.entitlements удалены из активного проекта')
doc.add_paragraph('   • Нет com.apple.developer.networking.vpn.api в проекте')
doc.add_paragraph('   • Проверка project.pbxproj: 0 упоминаний VPN entitlements')
doc.add_paragraph('   • Нет VPN серверов в настройках')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('4. VPN API ВЫЗОВЫ ОТСУТСТВУЮТ:').bold = True
doc.add_paragraph('   • startVPNTunnel() - не выполняется (закомментирован и не компилируются)')
doc.add_paragraph('   • stopVPNTunnel() - не выполняется (закомментирован и не компилируются)')
doc.add_paragraph('   • Все VPN API методы закомментированы и не компилируются')
doc.add_paragraph('   • 0 активных VPN API вызовов в скомпилированном коде')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('5. ПРИЛОЖЕНИЕ ИСПОЛЬЗУЕТ ТОЛЬКО СТАНДАРТНЫЕ iOS API:').bold = True
doc.add_paragraph('   • URLSession для сетевых запросов')
doc.add_paragraph('   • Стандартные протоколы безопасности')
doc.add_paragraph('   • Нет системных расширений (network extensions)')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('6. ПРОВЕРКА ВСЕХ КОМПОНЕНТОВ IPA ФАЙЛА:').bold = True
doc.add_paragraph('   • Бинарный код — VPN отсутствует')
doc.add_paragraph('   • Resources — VPN отсутствует')
doc.add_paragraph('   • Info.plist — VPN отсутствует')
doc.add_paragraph('   • Entitlements — VPN отсутствует (только в BACKUPS, неактивные файлы)')
doc.add_paragraph('   • Frameworks — NetworkExtension отсутствует')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ПОЛНОЕ УДАЛЕНИЕ VPN').bold = True
doc.add_paragraph('Все VPN-связанные компоненты были удалены:')
doc.add_paragraph('• Удалены все VPN файлы и targets')
doc.add_paragraph('• Переименованы все VPN классы и методы')
doc.add_paragraph('• Обновлены все VPN строки в пользовательском интерфейсе')
doc.add_paragraph('• Удалены VPN entitlements и конфигурации')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('ПРОСЬБА К КОМАНДЕ APPLE:').bold = True
doc.add_paragraph('Для быстрого исправления проблемы, пожалуйста, укажите нам, где именно вы видите VPN-функциональность:')
doc.add_paragraph('1. Скриншот экрана приложения, где видна VPN функциональность')
doc.add_paragraph('2. Или точное указание места в коде/метаданных, где вы видите VPN')
doc.add_paragraph('3. Или название конкретного экрана/функции, которая выглядит как VPN')

doc.add_paragraph()
doc.add_paragraph('Мы провели полный аудит кода и не находим никакой VPN функциональности. Все проверки показывают, что код чист от VPN компонентов. Мы готовы оперативно исправить любое указанное вами место.')

doc.add_paragraph()
doc.add_paragraph('Если приложение по какой-то причине выглядит как VPN, но не является таковым, мы готовы переименовать любые элементы интерфейса или описания, которые могут вызывать такую ассоциацию.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Build 14 содержит все исправления.').bold = True

doc.add_paragraph()
doc.add_paragraph('С уважением,')
doc.add_paragraph('Сергей Хлыстов')
doc.add_paragraph('ALADDIN')
doc.add_paragraph('22 декабря 2025')

# Сохраняем документ
output_file = 'ОТВЕТ_APPLE_VPN_22_12_2025_BUILD_14.docx'
output_path = os.path.join(os.path.dirname(__file__), output_file)
doc.save(output_path)

print(f"✅ DOCX файл создан: {output_path}")
print(f"📄 Файл содержит: английский текст + разрыв страницы + русский текст")
