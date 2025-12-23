#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для создания DOCX файла с ответом Apple по корпоративному аккаунту
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_docx():
    doc = Document()
    
    # Путь к файлам
    script_dir = os.path.dirname(os.path.abspath(__file__))
    russian_file = os.path.join(script_dir, 'ОТВЕТ_APPLE_КОРПОРАТИВНЫЙ_АККАУНТ_19_12_2025.txt')
    english_file = os.path.join(script_dir, 'ОТВЕТ_APPLE_КОРПОРАТИВНЫЙ_АККАУНТ_19_12_2025_ENGLISH.txt')
    
    # Читаем русскую версию
    with open(russian_file, 'r', encoding='utf-8') as f:
        russian_text = f.read()
    
    # Добавляем русскую версию
    for line in russian_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            # Делаем заголовки жирными
            if (line.strip().isupper() and ':' in line) or line.strip().startswith('ALADDIN') or line.strip().startswith('ОТВЕТ') or line.strip().startswith('Build'):
                if len(p.runs) > 0:
                    p.runs[0].bold = True
        else:
            doc.add_paragraph()
    
    # Разрыв страницы
    doc.add_page_break()
    
    # Читаем английскую версию
    with open(english_file, 'r', encoding='utf-8') as f:
        english_text = f.read()
    
    # Добавляем английскую версию
    for line in english_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            # Делаем заголовки жирными
            if (line.strip().isupper() and ':' in line) or line.strip().startswith('ALADDIN') or line.strip().startswith('RESPONSE') or line.strip().startswith('Build'):
                if len(p.runs) > 0:
                    p.runs[0].bold = True
        else:
            doc.add_paragraph()
    
    # Сохраняем
    output_path = '/Users/sergejhlystov/Downloads/ОТВЕТ_APPLE_КОРПОРАТИВНЫЙ_АККАУНТ_19_12_2025.docx'
    doc.save(output_path)
    print(f'✅ DOCX файл создан: {output_path}')

if __name__ == '__main__':
    create_docx()
