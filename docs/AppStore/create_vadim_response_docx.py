#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для создания DOCX файла с письмом для Вадима из Apple
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def process_content(doc, content, is_english=False):
    """Обрабатывает содержимое и добавляет в документ"""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Заголовок
        if line.startswith('Good day') or line.startswith('Добрый день'):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            continue
        
        # Обычный текст
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

def create_docx():
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Путь к файлам
    script_dir = os.path.dirname(os.path.abspath(__file__))
    english_file = os.path.join(script_dir, 'ОТВЕТ_ВАДИМУ_APPLE_ENGLISH_22_12_2025.txt')
    russian_file = os.path.join(script_dir, 'ОТВЕТ_ВАДИМУ_APPLE_RUSSIAN_22_12_2025.txt')
    
    # Добавляем заголовок "ENGLISH VERSION"
    p = doc.add_paragraph("ENGLISH VERSION")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    doc.add_paragraph()
    
    # Читаем и добавляем английскую версию
    with open(english_file, 'r', encoding='utf-8') as f:
        english_content = f.read()
    
    process_content(doc, english_content, is_english=True)
    
    # Разрыв страницы
    doc.add_page_break()
    
    # Добавляем заголовок "RUSSIAN VERSION"
    p = doc.add_paragraph("RUSSIAN VERSION / РУССКАЯ ВЕРСИЯ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    doc.add_paragraph()
    
    # Читаем и добавляем русскую версию
    with open(russian_file, 'r', encoding='utf-8') as f:
        russian_content = f.read()
    
    process_content(doc, russian_content, is_english=False)
    
    # Сохраняем в папку Загрузки
    downloads_path = os.path.expanduser('~/Downloads')
    output_file = os.path.join(downloads_path, 'ОТВЕТ_ВАДИМУ_APPLE_22_12_2025.docx')
    
    doc.save(output_file)
    print(f"✅ Файл создан: {output_file}")
    return output_file

if __name__ == '__main__':
    create_docx()

