#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для создания DOCX файла с ответом Apple по Guideline 2.1
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_horizontal_line(paragraph):
    """Добавляет горизонтальную линию"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr, 
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def process_content(doc, content, is_english=False):
    """Обрабатывает содержимое и добавляет в документ"""
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Заголовок письма
        if line.startswith('ОТВЕТ APPLE') or line.startswith('RESPONSE TO APPLE'):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
            continue
        
        # Дата
        if 'декабря 2025' in line or 'December 22, 2025' in line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)
            continue
        
        # Приветствие
        if line.startswith('Уважаемая команда') or line.startswith('Dear'):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)
            continue
        
        # Разделители
        if line.startswith('━━'):
            p = doc.add_paragraph()
            add_horizontal_line(p)
            continue
        
        # Нумерованные заголовки (1., 2., 3. и т.д.)
        if line and line[0].isdigit() and ('.' in line[:3] or line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.') or line.startswith('5.') or line.startswith('6.') or line.startswith('7.')):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.bold = True
            continue
        
        # Подзаголовки с маркерами
        if line.startswith('•') or line.startswith('-'):
            p = doc.add_paragraph(line, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(11)
            continue
        
        # Итоговая сводка
        if line.startswith('✅') or line.startswith('❌') or line.startswith('ИТОГОВАЯ СВОДКА') or line.startswith('SUMMARY'):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)
                if 'ИТОГОВАЯ' in line or 'SUMMARY' in line:
                    run.font.bold = True
            continue
        
        # Контакты и подпись
        if line.startswith('Контакты:') or line.startswith('Contacts:') or line.startswith('С уважением') or line.startswith('Best regards'):
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(11)
            continue
        
        # Обычный текст
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

def create_docx():
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Путь к файлам
    script_dir = os.path.dirname(os.path.abspath(__file__))
    english_file = os.path.join(script_dir, 'ОТВЕТ_APPLE_GUIDELINE_2.1_ENGLISH_22_12_2025.txt')
    russian_file = os.path.join(script_dir, 'ОТВЕТ_APPLE_GUIDELINE_2.1_КРАТКАЯ_ВЕРСИЯ_22_12_2025.txt')
    
    # Добавляем заголовок "ENGLISH VERSION"
    p = doc.add_paragraph("ENGLISH VERSION")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    doc.add_paragraph()
    
    # Читаем и добавляем английскую версию
    with open(english_file, 'r', encoding='utf-8') as f:
        english_content = f.read()
    
    process_content(doc, english_content, is_english=True)
    
    # Разрыв страницы
    doc.add_page_break()
    
    # Добавляем заголовок "RUSSIAN VERSION"
    p = doc.add_paragraph("RUSSIAN VERSION / РУССКАЯ ВЕРСИЯ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    doc.add_paragraph()
    
    # Читаем и добавляем русскую версию
    with open(russian_file, 'r', encoding='utf-8') as f:
        russian_content = f.read()
    
    process_content(doc, russian_content, is_english=False)
    
    # Сохраняем в папку Загрузки
    downloads_path = os.path.expanduser('~/Downloads')
    output_file = os.path.join(downloads_path, 'APPLE_GUIDELINE_2.1_22_12_2025.docx')
    
    doc.save(output_file)
    print(f"✅ Файл создан: {output_file}")
    return output_file

if __name__ == '__main__':
    create_docx()

